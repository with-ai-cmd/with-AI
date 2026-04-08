from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

doc = Document()

# ページ設定（A4）
section = doc.sections[0]
section.page_width = Cm(21.0)
section.page_height = Cm(29.7)
section.top_margin = Cm(2.5)
section.bottom_margin = Cm(2.5)
section.left_margin = Cm(2.5)
section.right_margin = Cm(2.5)

# デフォルトフォント設定
style = doc.styles['Normal']
font = style.font
font.name = 'Yu Mincho'
font.size = Pt(10.5)
rFonts = style.element.rPr.rFonts if style.element.rPr is not None else style.element.get_or_add_rPr().get_or_add_rFonts()
rFonts.set(qn('w:eastAsia'), 'Yu Mincho')

def set_run_font(run, bold=False, size=None):
    run.font.name = 'Yu Mincho'
    r = run._element
    rPr = r.get_or_add_rPr()
    rFonts = rPr.get_or_add_rFonts()
    rFonts.set(qn('w:eastAsia'), 'Yu Mincho')
    if bold:
        run.bold = True
    if size:
        run.font.size = size

def add_title(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.space_after = Pt(24)
    p.space_before = Pt(12)
    run = p.add_run(text)
    set_run_font(run, bold=True, size=Pt(16))

def add_article_title(text):
    p = doc.add_paragraph()
    p.space_before = Pt(18)
    p.space_after = Pt(6)
    run = p.add_run(text)
    set_run_font(run, bold=True, size=Pt(11))

def add_body(text):
    p = doc.add_paragraph()
    p.space_before = Pt(2)
    p.space_after = Pt(2)
    pf = p.paragraph_format
    pf.first_line_indent = Cm(0)
    run = p.add_run(text)
    set_run_font(run, size=Pt(10.5))
    return p

def add_numbered(num, text):
    p = doc.add_paragraph()
    p.space_before = Pt(2)
    p.space_after = Pt(2)
    pf = p.paragraph_format
    pf.left_indent = Cm(0.5)
    pf.first_line_indent = Cm(-0.5)
    run = p.add_run(f'{num}. {text}')
    set_run_font(run, size=Pt(10.5))
    return p

def add_bullet(text, indent=1.0):
    p = doc.add_paragraph()
    p.space_before = Pt(1)
    p.space_after = Pt(1)
    pf = p.paragraph_format
    pf.left_indent = Cm(indent)
    pf.first_line_indent = Cm(-0.3)
    run = p.add_run(f'・{text}')
    set_run_font(run, size=Pt(10.5))
    return p

def add_sub_bullet(text, indent=1.5):
    p = doc.add_paragraph()
    p.space_before = Pt(1)
    p.space_after = Pt(1)
    pf = p.paragraph_format
    pf.left_indent = Cm(indent)
    pf.first_line_indent = Cm(-0.3)
    run = p.add_run(f'- {text}')
    set_run_font(run, size=Pt(10))
    return p

def add_table(rows_data):
    table = doc.add_table(rows=len(rows_data), cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'
    for i, (col1, col2) in enumerate(rows_data):
        cell0 = table.cell(i, 0)
        cell1 = table.cell(i, 1)
        cell0.width = Cm(3.5)
        cell1.width = Cm(12.5)
        p0 = cell0.paragraphs[0]
        run0 = p0.add_run(col1)
        set_run_font(run0, bold=True, size=Pt(10))
        p1 = cell1.paragraphs[0]
        run1 = p1.add_run(col2)
        set_run_font(run1, size=Pt(10))
    return table

def add_empty_line(count=1):
    for _ in range(count):
        p = doc.add_paragraph()
        p.space_before = Pt(0)
        p.space_after = Pt(0)
        run = p.add_run('')
        run.font.size = Pt(6)

def add_signature_line(label, items):
    p = doc.add_paragraph()
    p.space_before = Pt(12)
    p.space_after = Pt(6)
    run = p.add_run(label)
    set_run_font(run, bold=True, size=Pt(11))
    for key, val in items:
        p2 = doc.add_paragraph()
        p2.space_before = Pt(2)
        p2.space_after = Pt(2)
        pf = p2.paragraph_format
        pf.left_indent = Cm(0.5)
        run2 = p2.add_run(f'{key}：{val}')
        set_run_font(run2, size=Pt(10.5))
    p3 = doc.add_paragraph()
    p3.space_before = Pt(8)
    pf3 = p3.paragraph_format
    pf3.left_indent = Cm(0.5)
    run3 = p3.add_run('印：＿＿＿＿＿＿＿＿')
    set_run_font(run3, size=Pt(10.5))


# === 契約書本文 ===

add_title('AIKOMON（社外AI責任者）業務委託契約書')

add_body('with-AI株式会社（以下「甲」という。）と、＿＿＿＿＿＿＿＿（以下「乙」という。）は、甲が乙に対して社外AI責任者サービス「AIKOMON」を提供するにあたり、以下のとおり業務委託契約（以下「本契約」という。）を締結する。')

# 第1条
add_article_title('第1条（目的）')
add_body('本契約は、甲が乙に対し、乙の事業におけるAI活用の戦略立案、業務再設計、AI実装・自動化開発、社内定着支援等を包括的に行う社外AI責任者サービス「AIKOMON」（以下「本サービス」という。）を提供することを目的とする。')

# 第2条（①変更なし）
add_article_title('第2条（業務内容）')
add_numbered(1, '甲は乙に対し、AI相談プランとして以下のサービスを提供する。')
add_empty_line()
add_table([
    ('月額料金', '49,800円（税別）'),
    ('打ち合わせ', '月1回のオンラインミーティング'),
    ('相談対応', 'AI活用に関する相談対応'),
    ('提案', '改善テーマの提案・優先順位整理'),
    ('業務再設計', '方向性の立案'),
    ('備考', '実務作業（手を動かす作業）・AI研修は含まない'),
])
add_empty_line()
add_numbered(2, '上記に含まれない追加業務が発生する場合は、甲乙間で別途協議のうえ、追加費用および内容を書面にて合意するものとする。')

# 第3条（②修正：3ヶ月契約＋自動更新）
add_article_title('第3条（契約期間および解約）')
add_numbered(1, '本契約の有効期間は、甲乙間で合意した開始日から3ヶ月間とし、期間満了日の1か月前までにいずれの当事者からも何らの意思表示なき場合、同じ条件でさらに3ヶ月間更新され、その後も同様とする。')
add_numbered(2, '甲または乙は、理由の如何を問わず、解約を希望する月の1ヶ月前までに、相手方に書面（電子メールを含む。）にて通知することにより、本契約を解約することができる。')
add_numbered(3, '解約の効力は、解約通知が相手方に到達した日の翌月末日をもって生じるものとする。')

# 第4条
add_article_title('第4条（委託料および支払条件）')
add_numbered(1, '乙は甲に対し、月額49,800円（税別）に法定の消費税額を加算した金額を支払うものとする。')
add_numbered(2, '契約開始月において、サービス提供期間が1ヶ月に満たない場合は、当該月の暦日数を基準に日割り計算した金額を委託料とする。（計算式：月額料金 ÷ 当該月の暦日数 × サービス提供日数）')
add_numbered(3, '甲は乙に対し、毎月末日締めにて請求書を発行し、乙は翌月末日までに甲の指定する銀行口座に振込にて支払うものとする。振込手数料は乙の負担とする。')
add_numbered(4, '追加業務が発生した場合の費用は、別途見積書に基づき甲乙間で合意するものとする。')

# 第5条
add_article_title('第5条（甲の義務）')
add_numbered(1, '甲は、善良なる管理者の注意をもって本サービスを提供するものとする。')
add_numbered(2, '甲は、AI技術の最新動向を継続的にキャッチアップし、乙の事業に最適なAI活用の提案・実行を行うものとする。')
add_numbered(3, '甲は、本サービスの遂行にあたり、乙の事業内容・業務フロー・経営方針を理解し、乙の経営目的に沿ったAI活用を設計するものとする。')
add_numbered(4, '甲は、第2条に定める頻度・範囲で打ち合わせおよび業務遂行を行うものとする。')

# 第6条
add_article_title('第6条（乙の協力義務）')
add_numbered(1, '乙は、甲が本サービスを円滑に遂行するために必要な情報、資料、データ等を適時に提供するものとする。')
add_numbered(2, '乙は、甲の合理的な提案・依頼に対し、社内の関係部署との調整や意思決定を適時に行うよう努めるものとする。')
add_numbered(3, '乙が前各項の協力義務を怠ったことにより、本サービスの遂行に遅延または支障が生じた場合、甲はその責任を負わないものとする。')

# 第7条
add_article_title('第7条（業務遂行体制）')
add_numbered(1, '甲は、本サービスの主たる担当者（以下「AI責任者」という。）を選任し、乙に通知するものとする。')
add_numbered(2, 'AI責任者の変更が必要な場合、甲は事前に乙に通知し、引継ぎを十分に行うものとする。')
add_numbered(3, '乙は、本サービスに関する社内の窓口担当者を選任し、甲に通知するものとする。')

# 第8条
add_article_title('第8条（成果物の取扱い）')
add_numbered(1, '本サービスの遂行により甲が作成した成果物（業務フロー設計書、AI実装物、自動化ツール、研修資料等）に関する著作権（著作権法第27条および第28条に規定する権利を含む。）は、以下のとおり取り扱うものとする。')
add_bullet('乙専用に作成された成果物（乙の業務フローに基づくAI実装物、自動化ツール、業務設計書等）：月額料金の完済をもって乙に帰属する')
add_bullet('甲の汎用的なノウハウ・手法・フレームワーク（AI活用の方法論、研修カリキュラムの構成、汎用ツール等）：甲に帰属する')
add_numbered(2, '甲は、乙への本サービス提供を通じて得た汎用的な知見・ノウハウを、乙の秘密情報を含まない形で、他の顧客へのサービス提供に活用することができるものとする。')
add_numbered(3, '甲が本サービスにおいて使用する第三者のソフトウェア、APIサービス等の知的財産権は、当該第三者に帰属する。乙がこれらを利用する場合、当該第三者の利用規約に従うものとする。')

# 第9条
add_article_title('第9条（秘密保持）')
add_numbered(1, '甲および乙は、本契約の締結および履行に関連して知り得た相手方の技術上、営業上その他一切の秘密情報（以下「秘密情報」という。）を、相手方の事前の書面による承諾なく、第三者に開示または漏洩してはならない。')
add_numbered(2, '秘密情報には、以下の情報を含むものとする。')
add_bullet('乙の経営情報、事業戦略、顧客情報、取引先情報、財務情報')
add_bullet('乙の業務フロー、業務データ、社内ノウハウ')
add_bullet('甲のAI技術、独自ノウハウ、アルゴリズム、システム設計情報')
add_bullet('本サービスの遂行過程で作成された資料、データ、成果物の内容')
add_bullet('本契約の存在および内容')
add_numbered(3, 'ただし、以下のいずれかに該当する情報は秘密情報から除くものとする。')
add_bullet('開示時点で既に公知であった情報')
add_bullet('開示後、受領者の責めに帰さない事由により公知となった情報')
add_bullet('開示時点で既に受領者が正当に保有していた情報')
add_bullet('正当な権限を有する第三者から秘密保持義務を負うことなく取得した情報')
add_bullet('秘密情報によらず独自に開発した情報')
add_numbered(4, '本条の義務は、本契約終了後も5年間存続するものとする。')

# 第10条
add_article_title('第10条（個人情報の取扱い）')
add_numbered(1, '甲は、本サービスの遂行にあたり乙から取扱いを委託された個人情報（個人情報の保護に関する法律に定義するものをいう。）について、善良なる管理者の注意をもって管理し、本サービスの目的以外に利用してはならない。')
add_numbered(2, '甲は、乙から委託された個人情報について、漏洩、滅失、毀損等の事故が発生した場合、またはそのおそれが生じた場合には、直ちに乙に報告し、乙の指示に従い適切な措置を講じるものとする。')
add_numbered(3, '甲は、本サービスの終了後または乙の求めがあった場合には、乙から委託された個人情報を速やかに乙に返還し、または乙の指示に従い廃棄するものとする。')

# 第11条
add_article_title('第11条（AIツール・外部サービスの利用）')
add_numbered(1, '甲は、本サービスの遂行にあたり、ChatGPT、Claude、Gemini等の生成AIサービス、その他各種SaaS・APIサービスを利用する場合がある。')
add_numbered(2, '甲は、前項の外部サービスの利用にあたり、乙の秘密情報および個人情報の取扱いについて十分に注意し、当該外部サービスの利用規約およびプライバシーポリシーに従うものとする。')
add_numbered(3, '甲は、乙の秘密情報を外部のAIサービスに入力する必要がある場合、事前に乙の承諾を得るものとする。')
add_numbered(4, '外部サービスの仕様変更、サービス停止等により本サービスの内容に影響が生じる場合、甲は速やかに乙に通知し、代替手段を含め対応を協議するものとする。')

# 第12条
add_article_title('第12条（再委託）')
add_numbered(1, '甲は、乙の事前の書面による承諾なく、本サービスの全部または一部を第三者に再委託してはならない。')
add_numbered(2, '乙の承諾を得て再委託した場合であっても、甲は再委託先の行為について乙に対し一切の責任を負うものとする。')

# 第13条
add_article_title('第13条（非保証事項・免責）')
add_numbered(1, '甲は、本サービスの提供にあたり最善を尽くすものとするが、以下の事項について保証するものではない。')
add_bullet('AI活用による特定の売上増加、コスト削減等の具体的な経営成果')
add_bullet('提案したAIツール・手法が乙の全ての業務課題を解決すること')
add_bullet('AI技術の進歩・変化に伴うサービス内容の将来にわたる同一性')
add_numbered(2, '乙がAIツールを使用して行った経営判断、業務遂行の結果について、甲は責任を負わないものとする。最終的な意思決定および業務遂行の責任は乙に帰属する。')

# 第14条（③修正：第3項を削除）
add_article_title('第14条（損害賠償）')
add_numbered(1, '甲および乙は、本契約に違反して相手方に損害を与えた場合、相手方に対し、その損害を賠償する責任を負う。')
add_numbered(2, '損害賠償の範囲は、通常損害に限るものとし、逸失利益を含む特別損害については、予見の有無を問わず賠償責任を負わないものとする。ただし、故意または重大な過失による場合はこの限りでない。')
# 旧第3項（損害賠償の上限条項）は削除

# 第15条
add_article_title('第15条（契約の解除）')
add_numbered(1, '甲または乙は、相手方が以下の各号のいずれかに該当した場合、催告なしに直ちに本契約の全部または一部を解除することができる。')
add_bullet('本契約の条項に違反し、相当期間を定めた催告にもかかわらず是正されないとき')
add_bullet('手形・小切手の不渡りが発生したとき')
add_bullet('差押え、仮差押え、仮処分、強制執行または競売の申立てがあったとき')
add_bullet('破産手続開始、民事再生手続開始、会社更生手続開始、特別清算開始の申立てがあったとき')
add_bullet('解散（合併による場合を除く。）または事業の全部もしくは重要な一部の譲渡を決議したとき')
add_bullet('監督官庁から営業停止または営業許可・免許の取消処分を受けたとき')
add_bullet('その他、本契約を継続し難い重大な事由が生じたとき')
add_numbered(2, '前項に加え、第3条の定めに基づき、甲または乙はいつでも本契約を解約することができる。')

# 第16条
add_article_title('第16条（契約終了時の措置）')
add_numbered(1, '本契約が終了した場合、甲は以下の措置を講じるものとする。')
add_bullet('乙のために作成した成果物（業務設計書、AI実装物、自動化ツール等）を乙に引き渡すこと')
add_bullet('乙から提供を受けた資料、データ等を乙に返還し、または乙の指示に従い廃棄すること')
add_bullet('AI実装物の運用・保守に関する引継ぎ事項を書面にて乙に提供すること')
add_numbered(2, '甲は、契約終了後も乙からの合理的な問い合わせに対し、契約終了後1ヶ月間は誠意をもって対応するよう努めるものとする。')
add_numbered(3, '契約終了後の継続的な保守・運用サポートが必要な場合は、別途契約を締結するものとする。')

# 第17条
add_article_title('第17条（反社会的勢力の排除）')
add_numbered(1, '甲および乙は、それぞれ相手方に対し、本契約締結時および契約期間中において、自らならびにその役員および実質的に経営を支配する者が、暴力団、暴力団員、暴力団関係企業、総会屋、社会運動標ぼうゴロ、政治活動標ぼうゴロ、特殊知能暴力集団その他の反社会的勢力（以下「反社会的勢力」という。）に該当しないことを表明し、保証する。')
add_numbered(2, '甲および乙は、相手方が前項の表明・保証に違反した場合、催告なしに直ちに本契約を解除することができる。この場合、解除した当事者は、相手方に対し、損害賠償を請求することができる。')

# 第18条
add_article_title('第18条（権利義務の譲渡禁止）')
add_body('甲および乙は、相手方の事前の書面による承諾なく、本契約上の地位または本契約に基づく権利義務の全部もしくは一部を第三者に譲渡し、担保に供し、またはその他の処分をしてはならない。')

# 第19条
add_article_title('第19条（不可抗力）')
add_body('天災地変、戦争、テロ、暴動、感染症の蔓延、法令の改廃、政府機関の行為、電気通信回線の障害その他当事者の責めに帰すことのできない事由（以下「不可抗力」という。）により、本契約の全部または一部の履行が遅延し、または不能となった場合、いずれの当事者も相手方に対し責任を負わないものとする。ただし、不可抗力の影響を受けた当事者は、速やかに相手方にその旨を通知し、影響を最小限にするよう努めるものとする。')

# 第20条
add_article_title('第20条（存続条項）')
add_body('本契約が終了した場合であっても、第8条（成果物の取扱い）、第9条（秘密保持）、第10条（個人情報の取扱い）、第13条（非保証事項・免責）、第14条（損害賠償）、第16条（契約終了時の措置）、第18条（権利義務の譲渡禁止）および本条の規定は、その効力を存続するものとする。')

# 第21条
add_article_title('第21条（協議事項）')
add_body('本契約に定めのない事項または本契約の解釈について疑義が生じた場合は、甲乙誠意をもって協議し、円満に解決を図るものとする。')

# 第22条
add_article_title('第22条（合意管轄）')
add_body('本契約に関する一切の紛争については、東京地方裁判所を第一審の専属的合意管轄裁判所とする。')

# 締結文
add_empty_line(2)
add_body('本契約の成立を証するため、本書2通を作成し、甲乙記名押印のうえ、各1通を保有する。')

add_empty_line(2)

# 日付欄
p = doc.add_paragraph()
p.space_before = Pt(6)
run = p.add_run('契約締結日：＿＿＿＿年＿＿月＿＿日')
set_run_font(run, bold=True, size=Pt(10.5))

p = doc.add_paragraph()
p.space_before = Pt(4)
run = p.add_run('サービス開始日：＿＿＿＿年＿＿月＿＿日')
set_run_font(run, bold=True, size=Pt(10.5))

add_empty_line(2)

# 甲の署名欄
add_signature_line('甲（サービス提供者）', [
    ('会社名', 'with-AI株式会社'),
    ('所在地', '東京都渋谷区千駄ヶ谷5-16-10'),
    ('代表者', '代表取締役　勝又海斗'),
])

add_empty_line(2)

# 乙の署名欄
add_signature_line('乙（サービス利用者）', [
    ('会社名', '＿＿＿＿＿＿＿＿＿＿＿＿'),
    ('所在地', '＿＿＿＿＿＿＿＿＿＿＿＿'),
    ('代表者', '＿＿＿＿＿＿＿＿＿＿＿＿'),
])

add_empty_line(2)

# プラン確認
p = doc.add_paragraph()
p.space_before = Pt(12)
run = p.add_run('契約プラン：AI相談プラン（49,800円/月・税別）')
set_run_font(run, bold=True, size=Pt(11))

# 保存
output_path = '/Users/kaitomain/Desktop/AIKOMON_業務委託契約書_修正版.docx'
doc.save(output_path)
print(f'契約書を保存しました: {output_path}')
