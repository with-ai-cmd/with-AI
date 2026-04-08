from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import os

doc = Document()

# --- スタイル設定 ---
style = doc.styles['Normal']
font = style.font
font.name = 'Yu Gothic'
font.size = Pt(10.5)
font.color.rgb = RGBColor(0x33, 0x33, 0x33)
style.element.rPr.rFonts.set(qn('w:eastAsia'), 'Yu Gothic')

for section in doc.sections:
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

MAIN_COLOR = RGBColor(0x0A, 0x3B, 0x8E)
ACCENT_COLOR = RGBColor(0x48, 0xA8, 0xE1)
TEXT_COLOR = RGBColor(0x33, 0x33, 0x33)

def add_heading_styled(text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = MAIN_COLOR
        run.font.name = 'Yu Gothic'
        run.element.rPr.rFonts.set(qn('w:eastAsia'), 'Yu Gothic')
    return h

def add_para(text, bold=False, size=None, color=None, align=None, space_after=Pt(6)):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = 'Yu Gothic'
    run.element.rPr.rFonts.set(qn('w:eastAsia'), 'Yu Gothic')
    if bold:
        run.bold = True
    if size:
        run.font.size = size
    if color:
        run.font.color.rgb = color
    if align:
        p.alignment = align
    p.paragraph_format.space_after = space_after
    return p

def add_separator():
    p = doc.add_paragraph()
    run = p.add_run('━' * 50)
    run.font.color.rgb = ACCENT_COLOR
    run.font.size = Pt(8)
    run.font.name = 'Yu Gothic'
    run.element.rPr.rFonts.set(qn('w:eastAsia'), 'Yu Gothic')
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(12)

def set_cell_font(cell, text, bold=False, color=TEXT_COLOR, size=Pt(10)):
    cell.text = ''
    p = cell.paragraphs[0]
    run = p.add_run(text)
    run.font.name = 'Yu Gothic'
    run.element.rPr.rFonts.set(qn('w:eastAsia'), 'Yu Gothic')
    run.font.size = size
    run.font.color.rgb = color
    run.bold = bold

def set_cell_shading(cell, color_hex):
    shading = cell._element.get_or_add_tcPr()
    shading_elm = shading.makeelement(qn('w:shd'), {
        qn('w:val'): 'clear',
        qn('w:color'): 'auto',
        qn('w:fill'): color_hex
    })
    shading.append(shading_elm)

def add_bullet(text):
    p = doc.add_paragraph(style='List Bullet')
    run = p.add_run(text)
    run.font.name = 'Yu Gothic'
    run.element.rPr.rFonts.set(qn('w:eastAsia'), 'Yu Gothic')
    run.font.size = Pt(10.5)
    run.font.color.rgb = TEXT_COLOR

# =====================
# 表紙
# =====================
for _ in range(5):
    doc.add_paragraph()

add_para('株式会社タイム様', bold=True, size=Pt(14), color=ACCENT_COLOR, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=Pt(12))
add_para('AIを活用した営業力強化のご提案', bold=True, size=Pt(24), color=MAIN_COLOR, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=Pt(40))

add_separator()

add_para('2026年3月19日', size=Pt(11), align=WD_ALIGN_PARAGRAPH.CENTER, space_after=Pt(4))
add_para('with-AI株式会社', size=Pt(13), color=MAIN_COLOR, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=Pt(4), bold=True)
add_para('代表取締役　勝又海斗', size=Pt(11), align=WD_ALIGN_PARAGRAPH.CENTER, space_after=Pt(4))

doc.add_page_break()

# =====================
# ご挨拶
# =====================
add_heading_styled('ご挨拶', level=1)

add_para(
    '株式会社タイム様\n'
    '代表取締役社長様'
)

add_para(
    '拝啓　時下ますますご清栄のこととお慶び申し上げます。\n\n'
    'このたびは、近藤様を通じて貴社の営業現場における課題をお伺いする機会をいただき、'
    '誠にありがとうございました。\n\n'
    'お話を伺い、貴社の営業力をAIの活用によってさらに強化できる可能性を強く感じております。'
    '本資料では、私たちが考えるAI活用の方向性と、まずリスクなく始められるトライアルプランをご提案させていただきます。',
    space_after=Pt(16)
)

add_separator()

# =====================
# 貴社の課題認識
# =====================
add_heading_styled('お伺いした貴社の課題', level=1)

add_para(
    '近藤様とのお打ち合わせを通じて、以下の課題を認識いたしました。',
    space_after=Pt(12)
)

add_heading_styled('営業力にばらつきがある', level=2)
add_para(
    'ベテラン社員と若手社員の間で、顧客への提案の質に差が生じています。'
    '経験豊富な社員であれば対応できる案件でも、経験の浅い社員では機会を逃してしまうケースがあるとのことでした。'
)

add_heading_styled('提案後のフォローが漏れやすい', level=2)
add_para(
    'お客様にご提案をお送りした後、返信がないまま案件が流れてしまうことが多いとのことです。'
    '現在の成約率は25〜30%程度とのことですが、フォローの仕組みを整えることで、ここを引き上げる余地は大きいと考えます。'
)

add_heading_styled('過去のお客様へのアプローチが弱い', level=2)
add_para(
    'これまでWeb経由のお問い合わせを中心に成長されてきた分、'
    '過去にお取引のあったお客様への能動的なアプローチが少ない状況とのことでした。'
)

add_separator()

# =====================
# ご提案の方向性
# =====================
add_heading_styled('ご提案の方向性', level=1)

add_para(
    '私たちは、AIを「便利なツール」として導入するのではなく、'
    '営業の皆様一人ひとりの隣にアシスタントがいるような仕組みをつくることをご提案します。',
    size=Pt(11),
    space_after=Pt(16)
)

add_para('AIにできること（イメージ）', bold=True, size=Pt(11), color=MAIN_COLOR, space_after=Pt(8))

add_bullet('お客様の情報や過去のやり取りをもとに、最適な提案内容をAIがサポート')
add_bullet('フォローが必要な案件をAIが自動で検知し、対応漏れを防止')
add_bullet('過去のお客様への連絡のきっかけやネタをAIが提案し、アプローチしやすくする')

doc.add_paragraph()

# 注記ボックス
note_table = doc.add_table(rows=1, cols=1)
note_table.alignment = WD_TABLE_ALIGNMENT.CENTER
cell = note_table.cell(0, 0)
set_cell_shading(cell, 'EBF5FB')
cell.text = ''
p = cell.paragraphs[0]
run = p.add_run(
    '具体的にどのような仕組みをつくるかは、トライアル開始後に社員の皆様への\n'
    'ヒアリングを行い、現場に本当にフィットする形で設計いたします。'
)
run.font.name = 'Yu Gothic'
run.element.rPr.rFonts.set(qn('w:eastAsia'), 'Yu Gothic')
run.font.size = Pt(10)
run.font.color.rgb = MAIN_COLOR
run.bold = True

add_separator()

# =====================
# 進め方
# =====================
add_heading_styled('進め方のイメージ', level=1)

add_para(
    'いきなり大きな投資をするのではなく、まず小さく始めて成果を確認していただく形をとります。',
    space_after=Pt(12)
)

# ステップ形式
steps = [
    ('Step 1', '社員の皆様へのヒアリング', '現場で本当に困っていること、AIがあったら助かる場面を丁寧にお聞きします。'),
    ('Step 2', '課題の優先順位づけと施策の設計', 'ヒアリング結果をもとに、最も効果の大きい課題から着手する計画を立てます。'),
    ('Step 3', 'AIの仕組みづくりと実装', '計画に沿って、実際にAIを業務に組み込む仕組みをつくります。'),
    ('Step 4', '定着支援', '社員の皆様がAIを自然に使いこなせるよう、研修やサポートを行います。'),
]

step_table = doc.add_table(rows=len(steps), cols=2)
step_table.alignment = WD_TABLE_ALIGNMENT.CENTER

for i, (step, title, desc) in enumerate(steps):
    # 左セル：Step番号
    set_cell_font(step_table.cell(i, 0), step, bold=True, color=MAIN_COLOR, size=Pt(10))
    set_cell_shading(step_table.cell(i, 0), 'EBF5FB')
    step_table.cell(i, 0).width = Cm(2.5)

    # 右セル：内容
    cell = step_table.cell(i, 1)
    cell.text = ''
    p = cell.paragraphs[0]
    run_t = p.add_run(title + '\n')
    run_t.font.name = 'Yu Gothic'
    run_t.element.rPr.rFonts.set(qn('w:eastAsia'), 'Yu Gothic')
    run_t.font.size = Pt(10.5)
    run_t.font.color.rgb = TEXT_COLOR
    run_t.bold = True
    run_d = p.add_run(desc)
    run_d.font.name = 'Yu Gothic'
    run_d.element.rPr.rFonts.set(qn('w:eastAsia'), 'Yu Gothic')
    run_d.font.size = Pt(10)
    run_d.font.color.rgb = TEXT_COLOR
    step_table.cell(i, 1).width = Cm(13.5)

add_separator()

# =====================
# 費用
# =====================
doc.add_page_break()
add_heading_styled('トライアルプランのご案内', level=1)

add_para(
    'まずは3ヶ月間のトライアルとして、以下の条件でご提案いたします。',
    space_after=Pt(12)
)

table = doc.add_table(rows=4, cols=2)
table.alignment = WD_TABLE_ALIGNMENT.CENTER
table.style = 'Table Grid'

fee_data = [
    ('項目', '内容'),
    ('期間', '3ヶ月'),
    ('費用', '50万円（総額・税別）'),
    ('成果保証', '成果が出なかった場合は全額返金'),
]

for i, (col1, col2) in enumerate(fee_data):
    if i == 0:
        set_cell_font(table.cell(i, 0), col1, bold=True, color=RGBColor(0xFF, 0xFF, 0xFF))
        set_cell_font(table.cell(i, 1), col2, bold=True, color=RGBColor(0xFF, 0xFF, 0xFF))
        set_cell_shading(table.cell(i, 0), '0A3B8E')
        set_cell_shading(table.cell(i, 1), '0A3B8E')
    else:
        set_cell_font(table.cell(i, 0), col1, bold=True)
        set_cell_font(table.cell(i, 1), col2)
        if i % 2 == 0:
            set_cell_shading(table.cell(i, 0), 'F5F8FC')
            set_cell_shading(table.cell(i, 1), 'F5F8FC')

for row in table.rows:
    row.cells[0].width = Cm(4)
    row.cells[1].width = Cm(12)

doc.add_paragraph()

# 全額返金保証の強調
risk_table = doc.add_table(rows=1, cols=1)
risk_table.alignment = WD_TABLE_ALIGNMENT.CENTER
cell = risk_table.cell(0, 0)
set_cell_shading(cell, 'E8F5E9')
cell.text = ''
p = cell.paragraphs[0]
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('成果が出なければ全額返金。リスクなくお試しいただけます。')
run.font.name = 'Yu Gothic'
run.element.rPr.rFonts.set(qn('w:eastAsia'), 'Yu Gothic')
run.font.size = Pt(12)
run.font.color.rgb = RGBColor(0x2E, 0x7D, 0x32)
run.bold = True

doc.add_paragraph()

add_heading_styled('費用対効果について', level=2)

add_para(
    '貴社の1案件あたりの粗利は20〜30万円とお伺いしております。\n'
    '仮にAIの活用によって月に1〜2件でも成約が増えれば、トライアル費用は十分に回収できる計算です。\n\n'
    'また、提案業務の効率化やフォロー漏れの防止による間接的な効果も見込まれます。'
)

add_separator()

# =====================
# with-AIについて
# =====================
add_heading_styled('with-AI株式会社について', level=1)

info_table = doc.add_table(rows=5, cols=2)
info_table.alignment = WD_TABLE_ALIGNMENT.CENTER
info_table.style = 'Table Grid'

info_data = [
    ('項目', '内容'),
    ('会社名', 'with-AI株式会社'),
    ('代表', '勝又海斗'),
    ('事業内容', '社外AI責任者サービス「AIKOMON」を中心に、\n企業のAI導入・業務改善を支援'),
    ('特徴', '提案だけで終わらず、実際に手を動かして仕組みを作り、\n社員が使いこなせるようになるまで伴走する実行型'),
]

for i, (col1, col2) in enumerate(info_data):
    if i == 0:
        set_cell_font(info_table.cell(i, 0), col1, bold=True, color=RGBColor(0xFF, 0xFF, 0xFF))
        set_cell_font(info_table.cell(i, 1), col2, bold=True, color=RGBColor(0xFF, 0xFF, 0xFF))
        set_cell_shading(info_table.cell(i, 0), '0A3B8E')
        set_cell_shading(info_table.cell(i, 1), '0A3B8E')
    else:
        set_cell_font(info_table.cell(i, 0), col1, bold=True)
        set_cell_font(info_table.cell(i, 1), col2)
        if i % 2 == 0:
            set_cell_shading(info_table.cell(i, 0), 'F5F8FC')
            set_cell_shading(info_table.cell(i, 1), 'F5F8FC')

for row in info_table.rows:
    row.cells[0].width = Cm(4)
    row.cells[1].width = Cm(12)

add_separator()

# =====================
# ネクストステップ
# =====================
add_heading_styled('ネクストステップ', level=1)

add_para(
    'ご興味をお持ちいただけましたら、まずは社長様を交えたお打ち合わせの場を設けさせていただければ幸いです。'
    '30分程度で、ご不明点やご懸念点にお答えいたします。',
    space_after=Pt(16)
)

next_table = doc.add_table(rows=4, cols=3)
next_table.alignment = WD_TABLE_ALIGNMENT.CENTER
next_table.style = 'Table Grid'

next_data = [
    ('', '内容', '時期'),
    ('1', 'お打ち合わせ（ご質問・ご懸念への回答）', '来週以降ご都合の良い日程'),
    ('2', 'トライアル契約の締結', '合意後'),
    ('3', '社員ヒアリング → トライアル開始', '契約後すぐ'),
]

for i, (col1, col2, col3) in enumerate(next_data):
    if i == 0:
        for j, val in enumerate([col1, col2, col3]):
            set_cell_font(next_table.cell(i, j), val, bold=True, color=RGBColor(0xFF, 0xFF, 0xFF))
            set_cell_shading(next_table.cell(i, j), '0A3B8E')
    else:
        set_cell_font(next_table.cell(i, 0), col1, bold=True, color=MAIN_COLOR)
        set_cell_font(next_table.cell(i, 1), col2)
        set_cell_font(next_table.cell(i, 2), col3)
        if i % 2 == 0:
            for j in range(3):
                set_cell_shading(next_table.cell(i, j), 'F5F8FC')

next_table.cell(0, 0).width = Cm(1.5)
next_table.cell(0, 1).width = Cm(10)
next_table.cell(0, 2).width = Cm(4.5)

doc.add_paragraph()
doc.add_paragraph()

add_para(
    'ご検討のほど、何卒よろしくお願い申し上げます。',
    align=WD_ALIGN_PARAGRAPH.CENTER,
    space_after=Pt(20)
)

add_para(
    'with-AI株式会社\n'
    '代表取締役　勝又海斗\n'
    'Email: info@with-ai.jp\n'
    'Tel: 090-1296-4814',
    align=WD_ALIGN_PARAGRAPH.RIGHT,
    size=Pt(10),
    color=MAIN_COLOR
)

# 保存
output_path = os.path.expanduser('~/Desktop/クロードコード/proposals/株式会社タイム_社内提案資料_20260319.docx')
os.makedirs(os.path.dirname(output_path), exist_ok=True)
doc.save(output_path)
print(f'Saved: {output_path}')
